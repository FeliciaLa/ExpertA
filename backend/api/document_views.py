from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from django.core.files.uploadedfile import UploadedFile
import mimetypes
import os
import PyPDF2
import docx
import io
from django.utils import timezone

from .models import Document, User
from .services import KnowledgeProcessor

class DocumentListView(APIView):
    """
    List all documents uploaded by the expert
    """
    permission_classes = [IsAuthenticated]
    
    def options(self, request, *args, **kwargs):
        # Handle CORS preflight requests - let Django CORS middleware handle headers
        return Response()
    
    def get(self, request):
        expert = request.user
        documents = Document.objects.filter(expert=expert)
        
        result = []
        for doc in documents:
            result.append({
                'id': doc.id,
                'filename': doc.filename,
                'file_size': doc.file_size,
                'mime_type': doc.mime_type,
                'upload_date': doc.upload_date,
                'status': doc.status,
                'error_message': doc.error_message
            })
        
        return Response({
            'documents': result,
            'total': len(result)
        })


class DocumentUploadView(APIView):
    """
    Upload documents for AI training
    """
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]
    
    def options(self, request, *args, **kwargs):
        # Handle CORS preflight requests with explicit headers
        response = Response()
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS"
        response["Access-Control-Allow-Headers"] = "Content-Type, Authorization, X-Requested-With, Accept, Origin"
        response["Access-Control-Allow-Credentials"] = "true"
        return response
    
    def get(self, request):
        """Test endpoint to verify connectivity"""
        return Response({
            'message': 'Document upload endpoint is accessible',
            'user': str(request.user),
            'parsers': [str(p) for p in self.parser_classes]
        })
    
    def post(self, request):
        expert = request.user
        print(f"DocumentUploadView - request.FILES: {request.FILES}")
        print(f"DocumentUploadView - request.FILES.keys(): {list(request.FILES.keys())}")
        print(f"DocumentUploadView - request.data: {request.data}")
        print(f"DocumentUploadView - Content-Type: {request.content_type}")
        print(f"DocumentUploadView - Headers: {dict(request.headers)}")
        
        # Try to get files from different possible keys
        files = request.FILES.getlist('documents')
        if not files:
            # Try other possible field names
            for key in request.FILES.keys():
                print(f"DocumentUploadView - Found files under key: {key}")
                files = request.FILES.getlist(key)
                break
        
        if not files:
            print(f"DocumentUploadView - No files found in request.FILES")
            print(f"DocumentUploadView - Available keys: {list(request.FILES.keys())}")
            return Response({
                'error': 'No files provided'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_count = 0
        results = []
        
        for file in files:
            try:
                # Get file info
                filename = file.name
                file_size = file.size
                
                # Determine mime type
                mime_type, _ = mimetypes.guess_type(filename)
                if not mime_type:
                    mime_type = 'application/octet-stream'
                
                # Create document record
                document = Document.objects.create(
                    expert=expert,
                    file=file,
                    filename=filename,
                    file_size=file_size,
                    mime_type=mime_type
                )
                
                # Process document asynchronously (in a real app, this would be a background task)
                try:
                    self._process_document(document)
                    document.status = 'completed'
                    document.save()
                except Exception as e:
                    print(f"Error processing document: {str(e)}")
                    document.status = 'failed'
                    document.error_message = str(e)
                    document.save()
                
                results.append({
                    'id': document.id,
                    'filename': document.filename,
                    'status': document.status
                })
                
                uploaded_count += 1
                
            except Exception as e:
                print(f"Error uploading document: {str(e)}")
                results.append({
                    'filename': file.name if hasattr(file, 'name') else 'unknown',
                    'error': str(e)
                })
        
        response = Response({
            'message': f'Successfully uploaded {uploaded_count} documents',
            'uploaded_count': uploaded_count,
            'results': results
        })
        # Add explicit CORS headers to response
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Credentials"] = "true"
        return response
    
    def _process_document(self, document):
        """Process document content and add to knowledge base"""
        content = ""
        
        try:
            # Extract text based on file type
            if document.mime_type == 'text/plain':
                # Process text files - read from the file object directly
                content = document.file.read().decode('utf-8', errors='ignore')
                document.file.seek(0)  # Reset file pointer
            
            elif document.mime_type == 'application/pdf' or document.filename.lower().endswith('.pdf'):
                # Process PDF files
                content = self._extract_text_from_pdf_file(document.file)
            
            elif document.mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                                     'application/msword'] or document.filename.lower().endswith(('.docx', '.doc')):
                # Process Word documents
                content = self._extract_text_from_word_file(document.file)
            
            else:
                # For other file types, just note that we can't process them
                raise Exception(f"Unsupported file type: {document.mime_type}")
            
            # If we have content, process it asynchronously
            if content.strip():
                # Try to queue document processing to run asynchronously
                try:
                    from django_q.tasks import async_task
                    from .tasks import process_document_async
                    task_id = async_task(process_document_async, document.id, content)
                    print(f"Document {document.id} saved and queued for knowledge processing: {document.filename} (task: {task_id})")
                except Exception as async_error:
                    # If async processing fails, process synchronously
                    print(f"Async processing failed for document {document.id}, processing synchronously: {str(async_error)}")
                    try:
                        from .services import KnowledgeProcessor
                        processor = KnowledgeProcessor(document.expert)
                        processor.process_document(document.id, content)
                        document.knowledge_processed = True
                        document.knowledge_processed_at = timezone.now()
                        document.save()
                        print(f"Document {document.id} processed synchronously: {document.filename}")
                    except Exception as sync_error:
                        print(f"Both async and sync processing failed for document {document.id}: {str(sync_error)}")
                        # Continue anyway, the document is still uploaded
            else:
                raise Exception("No text content could be extracted from the document")
                
        except Exception as e:
            print(f"Error processing document {document.id}: {str(e)}")
            raise
    
    def _extract_text_from_pdf_file(self, file_obj):
        """Extract text from a PDF file object"""
        text = ""
        try:
            file_obj.seek(0)  # Ensure we're at the beginning
            pdf_reader = PyPDF2.PdfReader(file_obj)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            file_obj.seek(0)  # Reset file pointer
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_text_from_word_file(self, file_obj):
        """Extract text from a Word document file object"""
        try:
            file_obj.seek(0)  # Ensure we're at the beginning
            doc = docx.Document(file_obj)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            file_obj.seek(0)  # Reset file pointer
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error extracting text from Word document: {str(e)}")
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path):
        """Extract text from a PDF file (deprecated - kept for compatibility)"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_text_from_word(self, file_path):
        """Extract text from a Word document (deprecated - kept for compatibility)"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error extracting text from Word document: {str(e)}")
            raise Exception(f"Failed to extract text from Word document: {str(e)}")


class DocumentDeleteView(APIView):
    """
    Delete a document
    """
    permission_classes = [IsAuthenticated]
    
    def options(self, request, *args, **kwargs):
        # Handle CORS preflight requests - let Django CORS middleware handle headers
        return Response()
    
    def delete(self, request, document_id):
        expert = request.user
        
        try:
            document = Document.objects.get(id=document_id, expert=expert)
        except Document.DoesNotExist:
            return Response({
                'error': 'Document not found'
            }, status=status.HTTP_404_NOT_FOUND)
        
        document.delete()
        
        return Response({
            'message': 'Document deleted successfully'
        }) 